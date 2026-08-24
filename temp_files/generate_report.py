import os
import subprocess
import sys

def install_and_import(package):
    try:
        import docx
    except ImportError:
        print(f"{package} not found, installing...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
    finally:
        globals()["docx"] = __import__("docx")
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        globals()["Pt"] = Pt
        globals()["RGBColor"] = RGBColor
        globals()["WD_PARAGRAPH_ALIGNMENT"] = WD_PARAGRAPH_ALIGNMENT

install_and_import('python-docx')

doc = docx.Document()

# Title
title = doc.add_heading('BÁO CÁO KIỂM THỬ GIAO DIỆN VÀ LOGIC\nLUỒNG MUA VÉ MÁY BAY & QUẢN LÝ VÉ', level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Meta Info
doc.add_paragraph('Nền tảng: Website (https://transport-hub-dev.gotrust.vn/)', style='Normal')
doc.add_paragraph('Ngày kiểm thử: 28/04/2026', style='Normal')
doc.add_paragraph('')

# 1. Tóm tắt
h1 = doc.add_heading('1. Tóm tắt kết quả kiểm thử (Test Summary)', level=1)
doc.add_paragraph('Luồng kiểm thử: Tìm kiếm -> Thông tin hành khách -> Dịch vụ -> Thanh toán -> Quản lý Vé (Vé của tôi).', style='List Bullet')
p = doc.add_paragraph('Kết quả chung: ', style='List Bullet')
p.add_run('Hệ thống có nhiều lỗi nghiêm trọng về điều hướng (Blocker) trong luồng mua vé. Phần Quản lý vé hoạt động tốt về mặt truy xuất dữ liệu nhưng mắc lỗi UI/UX sai lệch thông tin điểm đi/đến rất dễ gây hiểu lầm cho người dùng. Đặc biệt lỗi "Hệ thống phản hồi chậm" đã được xác minh là lỗi Timeout thực sự chứ không phải do Hết vé.').bold = True

doc.add_paragraph('')

# 2. Chi tiết
doc.add_heading('2. Chi tiết phát hiện (Detailed Findings)', level=1)

# Luồng mua vé
doc.add_heading('A. Luồng Đặt Vé & Thanh Toán', level=2)
p_bug0 = doc.add_paragraph('', style='List Bullet')
p_bug0.add_run('Lỗi Timeout "Hệ thống phản hồi chậm": ').bold = True
p_bug0.add_run('Đã kiểm chứng giả thuyết: Thông báo này KHÔNG PHẢI xuất hiện khi hết vé. Khi thực sự hết vé (VD: chặng ngách), hệ thống báo rất chuẩn "Rất tiếc hôm nay không còn vé nào...". Lỗi "Phản hồi chậm" là do hệ thống bị Timeout khi request API lấy dữ liệu chuyến bay (quá 15-20s). Đây là vấn đề lớn về Performance.')
doc.add_paragraph('Thông tin hành khách (Người đi cùng): Nhập thủ công và Chọn từ hồ sơ hoạt động trơn tru. Form validation bắt lỗi tốt.', style='List Bullet')
p_bug1 = doc.add_paragraph('', style='List Bullet')
p_bug1.add_run('Lỗi UI số điện thoại: ').bold = True
p_bug1.add_run('Hiển thị sai mã quốc gia (+66 hoặc +849 dù hiển thị cờ Việt Nam).')
p_bug2 = doc.add_paragraph('', style='List Bullet')
p_bug2.add_run('Nút "Tiếp tục" bị đơ & Navigation Loop (Blocker): ').bold = True
p_bug2.add_run('Thường xuyên không phản hồi hoặc load lại trang hiện tại khi chuyển từ Dịch vụ sang Thanh toán. Đơn hàng không được ghi nhận đẩy đủ.')
p_bug2.runs[0].font.color.rgb = RGBColor(255, 0, 0)
doc.add_paragraph('Trang Thanh toán (nếu ép truy cập): Tổng tiền hiển thị 0đ do không nhận được payload từ bước trước. Các liên kết footer (Ví, Hoạt động) bị lỗi 404.', style='List Bullet')

# Vé của tôi
doc.add_heading('B. Chức năng "Vé của tôi" (Quản lý vé)', level=2)
doc.add_paragraph('Truy cập: Vào từ trang chủ (đường dẫn /tickets). Danh mục hiển thị rõ ràng các tab Máy bay, Tàu hỏa, Xe khách, Metro. Tab "Sắp bay" và "Lịch sử" hoạt động ổn định.', style='List Bullet')
doc.add_paragraph('Chi tiết vé: Hoạt động tốt. Hiển thị chính xác Mã đặt chỗ (PNR), Hành khách, Hãng bay, v.v.', style='List Bullet')

doc.add_paragraph('Lỗi phát hiện:', style='Normal').bold = True
p_bug3 = doc.add_paragraph('', style='List Bullet')
p_bug3.add_run('Lỗi hoán đổi điểm Đi/Đến (Critical UI/Logic Bug): ').bold = True
p_bug3.add_run('Trên tất cả các thẻ tóm tắt vé máy bay, Tên thành phố ở tiêu đề bị ngược so với chặng bay thực tế ở chi tiết. Ví dụ: Tiêu đề ghi "Nha Trang -> Hà Nội", nhưng chuyến bay thực tế góc phải lại ghi "HAN -> CXR" (Hà Nội đi Cam Ranh). Lỗi này xuất hiện ở cả thẻ Sắp bay và Lịch sử, gây hiểu lầm nghiêm trọng cho khách hàng.')
p_bug3.runs[0].font.color.rgb = RGBColor(255, 0, 0)

p_bug4 = doc.add_paragraph('', style='List Bullet')
p_bug4.add_run('Thông báo trạng thái trống (Misleading Empty State): ').bold = True
p_bug4.add_run('Ở các tab không có vé (như Tàu hỏa, Metro), hệ thống báo "Không tìm thấy kết quả. Vui lòng thử lại từ khóa khác". Tuy nhiên trang này không hề có thanh tìm kiếm, câu thông báo này không hợp lý.')

doc.add_paragraph('')

# 3. Đề xuất
doc.add_heading('3. Đề xuất cải thiện', level=1)
doc.add_paragraph('Fix Timeout API: Tối ưu hóa thời gian phản hồi API tìm chuyến bay, tăng timeout limit hoặc dùng cơ chế caching/polling để tránh lỗi "Hệ thống phản hồi chậm".', style='List Number')
doc.add_paragraph('Fix Blocker Điều hướng: Sửa gấp lỗi vòng lặp trang khi bấm chuyển từ phần Dịch vụ sang Thanh toán.', style='List Number')
doc.add_paragraph('Fix Lỗi Hiển Thị Điểm Đi/Đến ở Vé của tôi: Đảo lại logic map tên thành phố trên thẻ tóm tắt vé sao cho khớp với mã sân bay chặng bay.', style='List Number')
doc.add_paragraph('Đồng bộ UI Tiền tố điện thoại & Thay đổi câu thông báo Empty State phù hợp hơn.', style='List Number')

doc.save('Bao_cao_test_ve_may_bay.docx')
print("File report created successfully.")
