import requests
import json
import docx
from docx.shared import Pt

# Token from user
TOKEN = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJuYW1laWQiOiI1MWJkOGMxOS0yZmMyLTQ2ZmYtY2FmMS0wOGQ5OWU5MTNhNzIiLCJhY3RvcnQiOiJOZ3V54buFbiBHaWFuZyBOYW0iLCJyb2xlIjoiVXNlciIsIlRpbWV6b25lT2Zmc2V0IjoiLTQyMCIsImp0aSI6ImNmZTVlY2U1LWE3OTgtNDQyOC04NDdhLTM2NTExYjI4NTI3MCIsIlBlcm1pc3Npb24iOiJyZXBvcnRNYmEudmlldyxyZXBvcnRNZXRlclRlc3QudmlldyxyZXBvcnRTZXR0bGVtZW50LnZpZXcscmVwb3J0MTc4MS52aWV3LHJlcG9ydEFtb3VudEVxdWlwbWVudC52aWV3LHJlcG9ydEpvYk9yZGVyU3RhdHVzLnZpZXcsd29ya1JlZ2lzdHJhdGlvbi5kZWxldGUsd29ya1JlZ2lzdHJhdGlvbi5yZWplY3Qsd29ya1JlZ2lzdHJhdGlvbi5hcHByb3ZlLHdvcmtSZWdpc3RyYXRpb24uaGlzdG9yeV9hcHByb3ZlZCxwZXJzb25hbFdvcmsudmlldyxwZXJzb25hbFdvcmsuY3JlYXRlLHBlcnNvbmFsV29yay51cGRhdGUscGVyc29uYWxXb3JrLmRlbGV0ZSxjb25maXJtU2hlZXRPZldvcmtMb2FkLnZpZXcsY29uZmlybVNoZWV0T2ZXb3JrTG9hZC51cGRhdGUsY29uZmlybVNoZWV0T2ZXb3JrTG9hZC5kZWxldGUsY29uZmlybVNoZWV0T2ZXb3JrTG9hZC5zZW5kX2NvbmZpcm0sZm9ybVJlcG9ydC52aWV3LGZvcm1SZXBvcnQuY3JlYXRlLGZvcm1SZXBvcnQudXBkYXRlLGZvcm1SZXBvcnQuZGVsZXRlLGZvcm1SZXBvcnQucHJpbnQsZm9ybVJlcG9ydC5zZW5kLGZvcm1SZXBvcnQuaGlzdG9yeV9"

HEADERS = {
    "Authorization": f"Bearer {TOKEN}",
    "Content-Type": "application/json"
}

doc = docx.Document()
doc.add_heading('Kết quả Test API QLTNKD (ĐKCT, BBKS, PATC)', 0)

def add_api_result(title, url, method, req_body, status_code, res_body):
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"URL: {url}")
    doc.add_paragraph(f"Method: {method}")
    if req_body:
        doc.add_paragraph(f"Request Body: {json.dumps(req_body, ensure_ascii=False)}")
    doc.add_paragraph(f"Status Code: {status_code}")
    
    try:
        pretty_res = json.dumps(res_body, indent=4, ensure_ascii=False)
        # truncate if too long
        if len(pretty_res) > 2000:
            pretty_res = pretty_res[:2000] + "\n... (truncated)"
        doc.add_paragraph(f"Response Body:\n{pretty_res}")
    except:
        res_str = str(res_body)
        if len(res_str) > 2000:
            res_str = res_str[:2000] + "\n... (truncated)"
        doc.add_paragraph(f"Response Body:\n{res_str}")

# 1. Test Paging
url_paging = "http://125.212.226.94:5006/api/workregistration/paging"
payload_paging = {"pageIndex": 1, "pageSize": 15}
print("Testing Paging API...")
try:
    res1 = requests.post(url_paging, headers=HEADERS, json=payload_paging)
    res1_json = res1.json()
    add_api_result("1. Lấy danh sách ĐKCT (Paging)", url_paging, "POST", payload_paging, res1.status_code, res1_json)
except Exception as e:
    res1_json = None
    add_api_result("1. Lấy danh sách ĐKCT (Paging)", url_paging, "POST", payload_paging, "ERROR", str(e))

dkct_id = None
if res1_json and isinstance(res1_json, dict) and res1_json.get("statusCode") == 200 and res1_json.get("data"):
    items = res1_json["data"].get("list", [])
    if items:
        dkct_id = items[0].get("id")

if dkct_id:
    # 2. Test Detail
    url_detail = f"http://125.212.226.94:5006/api/workregistration/{dkct_id}"
    print(f"Testing Detail API with ID: {dkct_id}...")
    try:
        res2 = requests.get(url_detail, headers=HEADERS)
        res2_json = res2.json()
        add_api_result("2. Lấy chi tiết ĐKCT", url_detail, "GET", None, res2.status_code, res2_json)
    except Exception as e:
        add_api_result("2. Lấy chi tiết ĐKCT", url_detail, "GET", None, "ERROR", str(e))

    # 3. Test PATC PDF
    url_patc = f"http://125.212.226.94:5006/api/workregistration/patc-pdf/{dkct_id}?isDownload=true"
    print("Testing PATC PDF API...")
    try:
        res3 = requests.get(url_patc, headers=HEADERS)
        add_api_result("3. API Lấy PDF Phương Án Thi Công", url_patc, "GET", None, res3.status_code, res3.text)
    except Exception as e:
        add_api_result("3. API Lấy PDF Phương Án Thi Công", url_patc, "GET", None, "ERROR", str(e))
else:
    doc.add_paragraph("Không có dữ liệu ĐKCT để test các API chi tiết.")

# 4. Try to get BBKS PDF (assume an arbitrary ID if we can't find one, or just test one from the detail)
bbks_id = None
if dkct_id and 'res2_json' in locals() and res2_json and isinstance(res2_json, dict):
    data = res2_json.get("data", {})
    bbks = data.get("bbksFiles", [])
    if bbks:
        bbks_id = bbks[0].get("surveyReportId")

if bbks_id:
    url_bbks = f"http://125.212.226.94:5006/api/surveyreport/{bbks_id}/pdf"
    print(f"Testing BBKS PDF API with ID: {bbks_id}...")
    try:
        res4 = requests.get(url_bbks, headers=HEADERS)
        add_api_result("4. API Lấy PDF Biên Bản Khảo Sát", url_bbks, "GET", None, res4.status_code, res4.text)
    except Exception as e:
        add_api_result("4. API Lấy PDF Biên Bản Khảo Sát", url_bbks, "GET", None, "ERROR", str(e))
else:
    doc.add_paragraph("Không tìm thấy BBKS ID để test API BBKS.")

doc.save('temp_files/Ket_Qua_Test_API.docx')
print("Done! Saved to Ket_Qua_Test_API.docx")
