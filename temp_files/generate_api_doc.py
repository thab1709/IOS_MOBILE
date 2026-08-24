import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_api_doc():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('HƯỚNG DẪN TEST API QLTNKD (ĐKCT, BBKS, PATC) TRÊN POSTMAN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Tài liệu này hướng dẫn cách gọi các API liên quan đến Đăng Ký Công Tác (ĐKCT), Biên Bản Khảo Sát (BBKS), và Phương Án Thi Công (PATC) trên Postman.')
    
    # Section 1
    doc.add_heading('1. Cấu hình xác thực (Authorization)', level=1)
    p = doc.add_paragraph('Tất cả các API dưới đây đều yêu cầu xác thực bằng Token. Trong Postman, chuyển sang tab ')
    p.add_run('Authorization').bold = True
    p.add_run(', chọn Type là ')
    p.add_run('Bearer Token').bold = True
    p.add_run(' và dán Token của bạn vào ô Token.')
    
    # Section 2
    doc.add_heading('2. Các API liên quan đến File PDF', level=1)
    
    doc.add_heading('2.1. Lấy PDF Biên Bản Khảo Sát (BBKS)', level=2)
    doc.add_paragraph('Method: GET\nURL: http://125.212.226.94:5006/api/surveyreport/{{id}}/pdf')
    doc.add_paragraph('Lưu ý: Thay {{id}} bằng ID của BBKS.')
    
    doc.add_heading('2.2. Lấy PDF Phương Án Thi Công (PATC)', level=2)
    doc.add_paragraph('Method: GET\nURL: http://125.212.226.94:5006/api/workregistration/patc-pdf/{{id}}?isDownload=true')
    doc.add_paragraph('Lưu ý: Thay {{id}} bằng ID của PATC (hoặc ID của phiếu ĐKCT liên quan).')
    
    # Section 3
    doc.add_heading('3. Các API Quản lý Đăng Ký Công Tác (ĐKCT)', level=1)
    
    doc.add_heading('3.1. Lấy danh sách ĐKCT (Có phân trang)', level=2)
    doc.add_paragraph('Method: POST\nURL: http://125.212.226.94:5006/api/workregistration/paging')
    doc.add_paragraph('Body (raw, JSON):')
    doc.add_paragraph('{\n    "pageIndex": 1,\n    "pageSize": 15\n}', style='Intense Quote')
    
    doc.add_heading('3.2. Lấy chi tiết ĐKCT', level=2)
    doc.add_paragraph('Method: GET\nURL: http://125.212.226.94:5006/api/workregistration/{{id}}')
    
    doc.add_heading('3.3. Gửi duyệt / Gửi xác nhận ĐKCT', level=2)
    doc.add_paragraph('Method: POST\nURL: http://125.212.226.94:5006/api/workregistration/send')
    doc.add_paragraph('Body (raw, JSON):')
    doc.add_paragraph('{\n    "ids": ["{{id}}"]\n}', style='Intense Quote')
    
    doc.add_heading('3.4. Xác nhận (Approve) ĐKCT', level=2)
    doc.add_paragraph('Method: POST\nURL: http://125.212.226.94:5006/api/workregistration/approve')
    doc.add_paragraph('Body (raw, JSON):')
    doc.add_paragraph('{\n    "ids": ["{{id}}"]\n}', style='Intense Quote')
    
    doc.add_heading('3.5. Từ chối (Reject) ĐKCT', level=2)
    doc.add_paragraph('Method: POST\nURL: http://125.212.226.94:5006/api/workregistration/reject')
    doc.add_paragraph('Body (raw, JSON):')
    doc.add_paragraph('{\n    "ids": ["{{id}}"],\n    "note": "Lý do từ chối"\n}', style='Intense Quote')
    
    doc.add_heading('3.6. Xem lịch sử ĐKCT', level=2)
    doc.add_paragraph('Method: GET\nURL: http://125.212.226.94:5006/api/workregistration/{{id}}/history')
    
    doc.add_heading('4. Ghi chú', level=1)
    doc.add_paragraph('Do môi trường yêu cầu token JWT thực tế (Bearer Token), bạn cần lấy token từ hệ thống đăng nhập (Console của Flutter app hoặc F12 trên web) để gắn vào Postman thì các request trên mới trả về dữ liệu thành công (Status 200).')

    doc.save('temp_files/Huong_Dan_Test_API_Postman.docx')

if __name__ == "__main__":
    create_api_doc()
